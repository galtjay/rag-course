import glob
from tqdm import tqdm
from PIL import Image
import chromadb
import os
import pdfplumber
import docx
import pandas as pd
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
from PIL.ExifTags import TAGS
import base64
import dashscope
from http import HTTPStatus
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

def extract_text_from_file(file_path):
    """
    根据文件类型提取文本内容。
    支持 .txt、.pdf、.docx、.epub、.csv、.xlsx 格式。
    """
    ext = file_path.split('.')[-1].lower()
    text = ''
    if ext == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
    elif ext == 'pdf':
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text()
    elif ext == 'docx':
        doc = docx.Document(file_path)
        # 提取正文内容
        for para in doc.paragraphs:
            text += para.text + '\n'
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + '\n'
    elif ext == 'epub':
        book = epub.read_epub(file_path)
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                # 使用BeautifulSoup解析HTML内容
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                text += soup.get_text() + '\n'
    elif ext == 'csv':
        df = pd.read_csv(file_path)
        text = df.to_string()
    elif ext == 'xlsx':
        df = pd.read_excel(file_path)
        text = df.to_string()
    return text


def chunk_text(text, chunk_size=500):
    """
    将文本分块，每个块的最大长度为 chunk_size。
    """
    chunks = []
    for i in range(0, len(text), chunk_size):
        chunks.append(text[i:i + chunk_size])
    print(chunks)
    return chunks


def process_files_in_directory(files):
    """
    遍历目录中的所有文件，提取文本并将其分块后通过API获取嵌入向量，并存储到chromadb。
    每个文件的元数据包含文件路径和原始文本内容。
    """

    # 初始化客户端
    client_openai = OpenAI(
        api_key=os.getenv("bailian_api_key"),  
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"  # 百炼服务的base_url
    )
    
    # 初始化ChromaDB客户端
    db_path = r"./vector_db"
    client_chroma = chromadb.PersistentClient(path=db_path)
    collection = client_chroma.get_or_create_collection(name='text2vec')
    
    # 遍历文件并处理
    for file in tqdm(files, desc="处理文件"):
        text = extract_text_from_file(file)
        print(f"处理文件：{file}, 提取文本长度：{len(text)}")  # 打印文件名和文本长度
        
        if text.strip():
            # 对每个文件进行文本分块
            file_chunks = chunk_text(text)
            
            # 将每个文件的文本块及其embeddings存入chromadb
            for idx, chunk in enumerate(file_chunks):
                try:
                    # 调用API获取文本嵌入向量
                    completion = client_openai.embeddings.create(
                        model="text-embedding-v3",
                        input=chunk,
                        dimensions=1024,
                        encoding_format="float"
                    )
                    
                    # 从API响应中提取嵌入向量
                    embedding = completion.data[0].embedding
                    
                    # 构建元数据
                    metadata = {
                        'file_path': file,  # 文件路径
                        'original_info': chunk  # 原始文本内容
                    }
                    print(metadata)
                    
                    # 构建唯一ID
                    ids = str(file) + "-" + str(idx)
                    
                    # 添加到集合
                    collection.add(
                        ids=ids,
                        metadatas=[metadata],
                        embeddings=[embedding]  # 确保embeddings是一个列表
                    )
                    print(ids, chunk)  # 用于调试
                    
                except Exception as e:
                    print(f"处理文本块出错 {file}-{idx}: {e}")
    
    print("所有文本块已成功存入数据库！")


def process_imgs_in_directory(imgs):
    """
    遍历指定文件夹中的图像，通过API生成嵌入并添加到集合。
    """
    
    db_path = r"./vector_db"
    client = chromadb.PersistentClient(path=db_path)
    
    # 获取或创建集合
    collection = client.get_or_create_collection(name='img2vec')
    
    # 将找到的图片向量化
    for image_path in tqdm(imgs, desc="图片向量化至数据库的进度"):
        try:
            # 打开图像并转换为RGB格式
            image = Image.open(image_path).convert("RGB")
            
            # 获取EXIF数据作为metadata
            exif_data = image.getexif()
            if not exif_data:
                exif_str = "No EXIF data found."
            else:
                exif_info = {}
                for tag_id, value in exif_data.items():
                    tag_name = TAGS.get(tag_id, tag_id)  # 获取EXIF标签名称
                    exif_info[tag_name] = value
                # 将EXIF信息转换为字符串
                exif_str = "\n".join([f"{key}: {value}" for key, value in exif_info.items()])
            
            # 确定图像格式
            img_format = image_path.split('.')[-1].lower()
            if img_format not in ['jpg', 'jpeg', 'png', 'bmp', 'gif', 'webp']:
                img_format = 'jpeg'  # 默认使用JPEG格式
            
            # 将图像转换为base64编码
            image_buffer = image.tobytes()
            with open(image_path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode('utf-8')
            
            # 设置图像格式
            image_data = f"data:image/{img_format};base64,{base64_image}"
            
            # 准备API输入
            inputs = [{'image': image_data}]
            dashscope.api_key = os.getenv("bailian_api_key")
            
            # 调用模型接口
            resp = dashscope.MultiModalEmbedding.call(
                model="multimodal-embedding-v1",
                input=inputs
            )
            
            if resp.status_code == HTTPStatus.OK:
                # 从API响应中提取嵌入向量
                embeddings = resp.output['embeddings'][0]['embedding']
                
                # 添加嵌入到集合
                metadata = {
                    'file_path': image_path,  # 文件路径
                    'original_info': exif_str  # 原始exif
                }
                print(metadata)
                
                collection.add(
                    ids=image_path,
                    metadatas=[metadata],
                    embeddings=[embeddings]  # 确保embeddings是一个列表
                )
            else:
                print(f"API调用失败 {image_path}: {resp.code}, {resp.message}")
                
        except Exception as e:
            print(f"处理图片出错 {image_path}: {e}")


def query_vec(query):
    """
    输入一段话，通过API获取嵌入向量，从chromadb数据库中找出最相近的文本和图片。
    """
    import os
    import chromadb
    import dashscope
    from openai import OpenAI
    from http import HTTPStatus
    
    # 初始化客户端和集合
    db_path = r"./vector_db"
    client = chromadb.PersistentClient(path=db_path)
    collection_text = client.get_or_create_collection(name='text2vec')
    collection_img = client.get_or_create_collection(name='img2vec')
    
    # 初始化OpenAI客户端(用于文本嵌入)
    client_openai = OpenAI(
        api_key=os.getenv("bailian_api_key"),
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
    )
    
    # 1. 获取文本嵌入并查询文本集合
    try:
        # 通过API获取文本嵌入
        completion = client_openai.embeddings.create(
            model="text-embedding-v3",
            input=query,
            dimensions=1024,
            encoding_format="float"
        )
        
        # 从API响应中提取嵌入向量
        query_embeddings_text = completion.data[0].embedding
        
        # 查询近似度前两个文本chunks
        results_text = collection_text.query(
            query_embeddings=[query_embeddings_text], 
            n_results=2, 
            include=["metadatas", "distances"]
        )
    except Exception as e:
        print(f"获取文本嵌入或查询出错: {e}")
        results_text = None
    
    # 2. 获取多模态嵌入并查询图片集合
    try:
        # 调用多模态API获取文本的图像嵌入表示
        dashscope.api_key = os.getenv("bailian_api_key")
        resp = dashscope.MultiModalEmbedding.call(
            model="multimodal-embedding-v1",
            input=[{'text': query}]  # 注意这里使用文本作为输入而不是图像
        )
        
        if resp.status_code == HTTPStatus.OK:
            # 从API响应中提取嵌入向量
            query_embeddings_img = resp.output['embeddings'][0]['embedding']
            
            # 查询近似度前两个图片
            results_img = collection_img.query(
                query_embeddings=[query_embeddings_img], 
                n_results=2, 
                include=["metadatas", "distances"]
            )
        else:
            print(f"多模态API调用失败: {resp.code}, {resp.message}")
            results_img = None
    except Exception as e:
        print(f"获取多模态嵌入或查询出错: {e}")
        results_img = None
    
    return [results_text, results_img]


def vector_all():
    directory = './files'
    finished_files_path = './vector_db/finished_files'
    # 获取所有文件路径
    all_files = glob.glob(os.path.join(directory, '**', '*'), recursive=True)
    all_files = [f for f in all_files if os.path.isfile(f)]
    print(all_files)
    # 读取已完成的文件列表
    if os.path.exists(finished_files_path):
        with open(finished_files_path, 'r', encoding='utf-8') as f:
            finished_files = set(f.read().splitlines())  # 读取已完成文件，存入集合（去重）
    else:
        finished_files = set()  # 如果文件不存在，则为空集合

    # 找出未完成的文件
    unfinished_files = [f for f in all_files if f not in finished_files]
    # 图片格式
    image_extensions = {'.png', '.jpg', '.jpeg', '.webp'}
    # 文本格式
    text_extensions = {'.txt', '.pdf', '.docx', '.epub', '.csv', '.xlsx'}

    # 分离未完成的图片和文本文件
    unfinished_imgs = [f for f in unfinished_files if os.path.splitext(f)[-1].lower() in image_extensions]
    unfinished_files = [f for f in unfinished_files if os.path.splitext(f)[-1].lower() in text_extensions]

    # 处理所有的文件
    process_files_in_directory(unfinished_files)
    process_imgs_in_directory(unfinished_imgs)

    # 写入新的 finished_files
    with open(finished_files_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(all_files))  # 更新为新的 all_files 列表


if __name__ == "__main__":
    vector_all()
    # print(query_vec("中国移动云手机"))